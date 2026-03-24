"""Unit tests for .docx text extraction (no full app import)."""

from io import BytesIO

import pytest
from docx import Document

from app.services.text_extraction import (
    DOCX_CONTENT_TYPE,
    extract_text_from_file,
    get_resolved_mime_type,
    upload_content_type_allowed,
)


def _minimal_docx_bytes() -> bytes:
    buf = BytesIO()
    doc = Document()
    doc.add_paragraph("Jane Doe")
    doc.add_paragraph("jane@example.com")
    doc.save(buf)
    return buf.getvalue()


def test_extract_text_from_docx_paragraphs():
    content = _minimal_docx_bytes()
    text = extract_text_from_file(content, DOCX_CONTENT_TYPE, "resume.docx")
    assert "Jane Doe" in text
    assert "jane@example.com" in text


def test_get_resolved_mime_type_docx_magic():
    content = _minimal_docx_bytes()
    assert get_resolved_mime_type(content, "application/pdf", None) == DOCX_CONTENT_TYPE


def test_upload_content_type_allowed_docx_filename_wrong_mime():
    assert upload_content_type_allowed("application/zip", "job.docx") is True


def test_upload_content_type_rejects_docx_filename_with_image_mime():
    assert upload_content_type_allowed("image/png", "cv.docx") is False


def test_non_docx_zip_rejected():
    # Minimal ZIP without word/document.xml
    import zipfile
    from io import BytesIO as BIO

    b = BIO()
    with zipfile.ZipFile(b, "w") as zf:
        zf.writestr("readme.txt", "x")
    content = b.getvalue()
    with pytest.raises(ValueError, match="Unsupported|not a valid DOCX"):
        extract_text_from_file(content, "application/zip", "not_word.zip")
