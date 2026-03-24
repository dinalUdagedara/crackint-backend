"""
Extract raw text from uploaded files: PDF (PyMuPDF), images (OCR via Tesseract),
and Word Open XML (.docx) via python-docx.
Used by resume and job extraction when the client uploads a file.
"""

import io
import zipfile
from typing import Optional

import pymupdf  # PyMuPDF (import as pymupdf, not fitz for clarity in deps)

# Word Open XML (.docx). Legacy binary .doc is not supported.
DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

# Content types accepted for file upload (PDF + images + DOCX). Use in routes for validation.
SUPPORTED_FILE_CONTENT_TYPES = frozenset({
    "application/pdf",
    "image/png",
    "image/jpeg",
    "image/jpg",
    "image/webp",
    DOCX_CONTENT_TYPE,
})

# User-facing detail for 400 responses (routes).
UNSUPPORTED_UPLOAD_DETAIL = (
    "Only PDF, Word (.docx), and images (PNG, JPEG, WebP) are accepted. "
    "Legacy .doc is not supported."
)


def _is_docx_ooxml(content: bytes) -> bool:
    """True if bytes are a ZIP package containing word/document.xml (OOXML Word)."""
    if len(content) < 4 or content[:2] != b"PK":
        return False
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as zf:
            return "word/document.xml" in zf.namelist()
    except zipfile.BadZipFile:
        return False


def _detect_content_type(content: bytes) -> Optional[str]:
    """Infer content type from magic bytes when client sends application/octet-stream."""
    if len(content) < 12:
        return None
    if content[:4] == b"%PDF":
        return "application/pdf"
    if _is_docx_ooxml(content):
        return DOCX_CONTENT_TYPE
    if content[:8] == b"\x89PNG\r\n\x1a\n":
        return "image/png"
    if content[:3] == b"\xff\xd8\xff":
        return "image/jpeg"
    if content[:4] == b"RIFF" and content[8:12] == b"WEBP":
        return "image/webp"
    return None


def resolve_effective_content_type(
    declared: Optional[str],
    filename: Optional[str],
    content: bytes,
) -> str:
    """
    Choose MIME for extraction: honor explicit supported types, use .docx filename
    when MIME is missing or clearly mislabeled (but not when claiming PDF/image).
    """
    d = (declared or "").strip().lower()
    fn = (filename or "").lower()

    # Prefer OOXML magic so mislabeled uploads (e.g. PDF MIME on a Word file) still work.
    if _is_docx_ooxml(content):
        return DOCX_CONTENT_TYPE

    if fn.endswith(".docx") and d not in ("application/pdf",) and not d.startswith("image/"):
        return DOCX_CONTENT_TYPE

    if d == DOCX_CONTENT_TYPE:
        return DOCX_CONTENT_TYPE

    if d in ("", "application/octet-stream"):
        detected = _detect_content_type(content)
        return detected or "application/octet-stream"

    return d


def get_resolved_mime_type(
    content: bytes,
    declared: Optional[str],
    filename: Optional[str],
) -> Optional[str]:
    """
    Final MIME after filename hints and magic-byte detection, or None if still unknown.
    Used by CV scoring to choose text (DOCX) vs vision (PDF/image).
    """
    ct = resolve_effective_content_type(declared, filename, content)
    if ct == "application/octet-stream":
        ct = _detect_content_type(content)
    if not ct or ct == "application/octet-stream":
        return None
    return ct


def upload_content_type_allowed(declared: Optional[str], filename: Optional[str]) -> bool:
    """
    Whether an upload may be attempted before reading bytes.
    Unknown / octet-stream is allowed; magic bytes decide later.
    """
    d = (declared or "").strip().lower()
    fn = (filename or "").lower()
    if fn.endswith(".docx") and (d.startswith("image/") or d == "application/pdf"):
        return False
    if not d or d == "application/octet-stream":
        return True
    if d in SUPPORTED_FILE_CONTENT_TYPES:
        return True
    if d in ("application/zip", "application/x-zip-compressed") and fn.endswith(".docx"):
        return True
    if fn.endswith(".docx"):
        return True
    return False


def _extract_text_from_pdf_via_ocr(file_content: bytes) -> str:
    """
    Extract text from PDF by rendering each page to an image and running Tesseract OCR.
    """
    from app.services.ocr import extract_text_from_image

    doc = pymupdf.open(stream=file_content, filetype="pdf")
    try:
        parts: list[str] = []
        for page in doc:
            # Render page at 200 DPI for reasonable OCR quality
            pix = page.get_pixmap(dpi=200, alpha=False)
            png_bytes = pix.tobytes("png")
            page_text = extract_text_from_image(png_bytes, content_type="image/png")
            if page_text.strip():
                parts.append(page_text.strip())
        return "\n\n".join(parts).strip() or ""
    finally:
        doc.close()


def _extract_text_from_pdf_direct(file_content: bytes) -> str:
    """Extract embedded text from PDF (no OCR). Returns empty if no text layer."""
    doc = pymupdf.open(stream=file_content, filetype="pdf")
    try:
        parts: list[str] = []
        for page in doc:
            text = page.get_text()
            if text:
                parts.append(text.strip())
        return "\n\n".join(parts).strip() or ""
    finally:
        doc.close()


def extract_text_from_pdf(file_content: bytes) -> str:
    """
    Extract plain text from PDF bytes.
    Always tries OCR first (render pages + Tesseract); if OCR returns empty, falls back to embedded text.

    Args:
        file_content: Raw PDF file bytes.

    Returns:
        Extracted text as a single string. Returns empty string if no text found.

    Raises:
        ValueError: If content is not valid PDF.
    """
    if not file_content or len(file_content) < 4:
        raise ValueError("Empty or invalid PDF content")

    # Always try OCR first; fall back to embedded text if OCR returns empty
    result = _extract_text_from_pdf_via_ocr(file_content)
    # Fall back to embedded text if OCR produced nothing
    if not result:
        result = _extract_text_from_pdf_direct(file_content)

    return result


def _extract_text_from_docx(file_content: bytes) -> str:
    """Extract plain text from .docx (paragraphs and table cells)."""
    from docx import Document

    if not _is_docx_ooxml(file_content):
        raise ValueError("File is not a valid DOCX (Office Open XML) document.")

    try:
        doc = Document(io.BytesIO(file_content))
    except Exception as e:
        raise ValueError(f"Invalid or corrupted DOCX file: {e}") from e

    parts: list[str] = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append("\t".join(cells))
    return "\n\n".join(parts).strip()


def extract_text_from_file(
    content: bytes,
    content_type: Optional[str] = None,
    filename: Optional[str] = None,
) -> str:
    """
    Extract plain text from file bytes (PDF, image, or DOCX). Dispatches by MIME
    (and optional filename for mislabeled DOCX).

    Args:
        content: Raw file bytes.
        content_type: MIME type from the client.
        filename: Original filename (optional); used to treat mislabeled .docx uploads.

    Returns:
        Extracted text as a single string.

    Raises:
        ValueError: If content is empty, type is unsupported, or extraction fails.
    """
    if not content or len(content) < 4:
        raise ValueError("Empty or invalid file content")

    ct = resolve_effective_content_type(content_type, filename, content)
    if ct == "application/octet-stream":
        ct = _detect_content_type(content)
    if not ct or ct == "application/octet-stream":
        raise ValueError(
            "Unsupported file type. Use PDF, Word (.docx), or image (PNG, JPEG, WebP)."
        )

    if ct == DOCX_CONTENT_TYPE:
        return _extract_text_from_docx(content)

    if ct == "application/pdf":
        return extract_text_from_pdf(content)

    if ct in SUPPORTED_FILE_CONTENT_TYPES and ct.startswith("image/"):
        from app.services.ocr import extract_text_from_image
        return extract_text_from_image(content, content_type=ct)

    raise ValueError(
        f"Unsupported file type: {content_type}. Use PDF, Word (.docx), or image (PNG, JPEG, WebP)."
    )
